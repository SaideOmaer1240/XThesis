from rest_framework import viewsets, permissions
from rest_framework.generics import DestroyAPIView
from rest_framework.decorators import action
from rest_framework.response import Response
from django.http import HttpResponse
from django.conf import settings
from docx import Document
from docx.enum.text import WD_BREAK
import os  
from .models import Thesis
from .serializers import ThesisSerializer
from .permissions import IsAuthor
from accounts.datetimes import DateTime 
import pypandoc
import re
class TopicViewSet(viewsets.ModelViewSet):
    serializer_class = ThesisSerializer
    permission_classes = [permissions.IsAuthenticated, IsAuthor]

    def get_queryset(self):         
        queryset = Thesis.objects.filter(author=self.request.user).order_by('-date_added')
        
        unique_titles = set()
        unique_queryset = []
     
        for thesis in queryset:
            if thesis.code not in unique_titles:
                unique_titles.add(thesis.code)
                unique_queryset.append(thesis)

        return unique_queryset

    def perform_create(self, serializer):
        serializer.save(author=self.request.user) 

class ThesisViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated, IsAuthor]
    serializer_class = ThesisSerializer

    def get_queryset(self):
        user = self.request.user
        code = self.request.query_params.get('code')
        queryset = Thesis.objects.filter(author=user)
        
        if code:
            queryset = queryset.filter(code=code)
        return queryset

    def perform_create(self, serializer):
        serializer.save(author=self.request.user)
        
            

    def copiar_estilos(self, origem, destino):
        destino.style = origem.style
        destino.paragraph_format.alignment = origem.paragraph_format.alignment
        destino.paragraph_format.left_indent = origem.paragraph_format.left_indent
        destino.paragraph_format.right_indent = origem.paragraph_format.right_indent
        destino.paragraph_format.first_line_indent = origem.paragraph_format.first_line_indent
        destino.paragraph_format.keep_together = origem.paragraph_format.keep_together
        destino.paragraph_format.keep_with_next = origem.paragraph_format.keep_with_next
        destino.paragraph_format.page_break_before = origem.paragraph_format.page_break_before
        destino.paragraph_format.widow_control = origem.paragraph_format.widow_control
        destino.paragraph_format.space_before = origem.paragraph_format.space_before
        destino.paragraph_format.space_after = origem.paragraph_format.space_after
        destino.paragraph_format.line_spacing = origem.paragraph_format.line_spacing
        destino.paragraph_format.line_spacing_rule = origem.paragraph_format.line_spacing_rule

    def localizar_paragrafo(self, document, texto):
        for paragraph in document.paragraphs:
            if texto in paragraph.text:
                return paragraph
        return None 

    def remover_paragrafo(self, doc, paragrafo):
        p = paragrafo._element
        p.getparent().remove(p)
        p._element = p._p = None

    def variable_content(self, institute, student, instructor, disciplina, topic, cidade):
        year = DateTime.year()
        mes = DateTime.get_month()
        credentials = {
            'ISaideOmar': institute,
            'CadeiroModuloDisciplina': disciplina,
            'TopicoTema': topic,
            'AutorDiscente': student,
            'SupervisorDocente': instructor,
            'CidadeAutorT': cidade,
            'MesAutorCria': mes,
            'AnoAutorCria': year
        }
        return credentials

    def replace_text(self, doc, replacements):
        def replace_in_paragraph(paragraph):
            for key, value in replacements.items():
                if key in paragraph.text:
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            inline[i].text = inline[i].text.replace(key, str(value))

        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.replace_text(cell, replacements)

    @action(detail=False, methods=['get'], url_path='gerar_documento')
    def gerar_documento(self, request):
        user = request.user
        code = request.query_params.get('code')

        if not code:
            return Response({'error': 'O parâmetro code é obrigatório.'}, status=400)

        first_thesis = self.get_queryset().filter(code=code).first()

        if first_thesis:
            cidade = first_thesis.cidade
            disciplina = first_thesis.disciplina
            institute = first_thesis.institute
            instructor = first_thesis.instructor
            student = first_thesis.student
        else:
            cidade = 'Quelimane'
            disciplina = 'Quimica'
            institute = 'Escola Secundária Geral de Quelimane'
            instructor = 'Antonio'
            student = 'Saíde Omar Saíde'

        theses = self.get_queryset().filter(code=code)
        if not theses.exists():
            return Response({'error': f'Tópico com código "{code}" não encontrado.'}, status=404)

        doc = Document(os.path.join(settings.MEDIA_ROOT, 'modelo.docx'))
        credentials = self.variable_content(
            cidade=cidade,
            disciplina=disciplina,
            institute=institute,
            instructor=instructor,
            student=student,
            topic=first_thesis.topic,
        )

        

        # Definir diretórios de documentos e arquivos temporários
        directory_doc = os.path.join(settings.MEDIA_ROOT, 'documents', user.username)
        directory_temp = os.path.join(settings.MEDIA_ROOT, 'tempFile', user.username)

        # Verificar se os diretórios existem; se não, criar
        os.makedirs(directory_doc, exist_ok=True)
        os.makedirs(directory_temp, exist_ok=True)

        # Remover arquivos antigos no diretório de documentos
        for root, dirs, files in os.walk(directory_doc):
            for file in files:
                if file.lower().endswith('.docx'):
                    os.remove(os.path.join(root, file))

        new_file_name = f'{theses[0].topic}.docx'
        file_path = os.path.join(directory_doc, new_file_name) 

        # Iterar sobre cada tese para criar o documento
        for thesis in theses:
            # Remover arquivos antigos no diretório temporário
            for root, dirs, files in os.walk(directory_temp):
                for file in files:
                    if file.lower().endswith('.docx'):
                        os.remove(os.path.join(root, file))
            temp_docx = 'temp_docx.docx'
            
            # Criar novo arquivo temporario
            temp_docx_path = os.path.join(directory_temp, temp_docx)
            latex = thesis.text

            # Converter LaTeX para DOCX
            pypandoc.convert_text(latex, 'docx', format='latex', outputfile=temp_docx_path)

            temp_doc = Document(temp_docx_path) 
            # Verificar se o arquivo final já existe e carregar o documento, caso contrário usar modelo
            if os.path.isfile(file_path):
                existing_doc = Document(file_path)
            else:
                existing_doc = Document(os.path.join(settings.MEDIA_ROOT, 'modelo.docx'))

            # Adicionar conteúdo do documento temporário ao documento existente
            for element in temp_doc.element.body:
                existing_doc.element.body.append(element)

            # Salvar o documento final
            existing_doc.save(file_path)

        # Reabrir o documento final para aplicar estilo aos parágrafos que iniciam com número seguido de ponto
        final_doc = Document(file_path)
        self.replace_text(final_doc, credentials)

        # Localizar um determinado paragrafo para copiar o seu estilo
        paragrafo_origem = self.localizar_paragrafo(final_doc, 'Copiar parágrafo')

        if paragrafo_origem:
            # Iterar sobre os parágrafos para encontrar o índice e aplicar alterações
            for index, para in enumerate(final_doc.paragraphs):
                # Verificar se o parágrafo corresponde ao padrão desejado
                if re.match(r'^\d+\.', para.text.strip()) or re.match(r'Introdução', para.text.strip()) or re.match(r'Referências Bibliográficas', para.text.strip()):
                    # Aplicar o estilo copiado ao parágrafo
                    self.copiar_estilos(paragrafo_origem, para)  # Substitua com a implementação correta da função
                    
                    if '. Revisão de Literatura' in para.text:
                        # Insere um novo parágrafo antes do parágrafo localizado
                        new_paragraph = final_doc.paragraphs[index].insert_paragraph_before()
                        # Adiciona uma quebra de página ao novo parágrafo
                        new_paragraph.add_run().add_break(WD_BREAK.PAGE)
                if re.match(r'Referências Bibliográficas', para.text.strip()): 
                    para.add_run().add_break(WD_BREAK.PAGE)
                    para.add_run().add_text("Referência Bibliográfica")
                
        
        # Procurar e remover a repetição de "Referência Bibliográfica"
        found = True
        for  i, paragraph in enumerate(final_doc.paragraphs): 
            if "Referências Bibliográficas" in paragraph.text:
                if found:
                    paragraph.text = paragraph.text.replace("Referências Bibliográficas", "")
                    found = False 
                new_paragraph = final_doc.paragraphs[i].insert_paragraph_before() 
                new_paragraph.add_run().add_break(WD_BREAK.PAGE) 
                
        self.remover_paragrafo(final_doc, paragrafo_origem)
        # Salvar o documento final após a aplicação dos estilos 
        final_doc.save(file_path)

        # Enviar o documento como resposta HTTP
        with open(file_path, 'rb') as arquivo_aberto:
            response = HttpResponse(
                arquivo_aberto.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{new_file_name}"'

        return response
        

class DestroyAllTheses(DestroyAPIView):
    permission_classes = [permissions.IsAuthenticated, IsAuthor]
    queryset = Thesis.objects.all()
    
    def get_object(self): 
        user = self.request.user
        return Thesis.objects.filter(author=user)

class DestroyOneThesis(DestroyAPIView):
    permission_classes = [permissions.IsAuthenticated, IsAuthor]
    
    def get_object(self):
        user = self.request.user
        code = self.request.query_params.get('code')
        queryset = Thesis.objects.filter(author=user)
        
        if code:
            queryset = queryset.filter(code=code)
            return queryset

