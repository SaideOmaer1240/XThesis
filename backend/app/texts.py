from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Função para copiar estilos de um parágrafo para outro
def copiar_estilos(origem, destino):
    destino.style = origem.style
    destino.alignment = origem.alignment
    destino.paragraph_format.left_indent = origem.paragraph_format.left_indent
    destino.paragraph_format.right_indent = origem.paragraph_format.right_indent
    destino.paragraph_format.first_line_indent = origem.paragraph_format.first_line_indent
    destino.paragraph_format.keep_together = origem.paragraph_format.keep_together
    destino.paragraph_format.keep_with_next = origem.paragraph_format.keep_with_next
    destino.paragraph_format.page_break_before = origem.paragraph_format.page_break_before
    destino.paragraph_format.widow_control = origem.paragraph_format.widow_control
    destino.paragraph_format.space_before = origem.paragraph_format.space_before
    destino.paragraph_format.space_after = origem.paragraph_format.space_after
    destino.paragraph_format.line_spacing = origem.paragraph_format.line_spacing
    destino.paragraph_format.line_spacing_rule = origem.paragraph_format.line_spacing_rule

# Função para localizar o parágrafo com um texto específico
def localizar_paragrafo_com_texto(doc, texto):
    for paragrafo in doc.paragraphs:
        if texto in paragrafo.text:
            return paragrafo
    return None

# Função para remover um parágrafo de um documento
def remover_paragrafo(doc, paragrafo):
    p = paragrafo._element
    p.getparent().remove(p)
    p._element = p._p = None

# Função para adicionar um campo de sumário automático
def adicionar_sumario(doc):
    # Cria o parágrafo onde o sumário será inserido
    paragrafo_sumario = doc.add_paragraph()
    
    # Cria o campo para o sumário
    run = paragrafo_sumario.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(fld_char3)

# Abrir o documento existente
doc = Document('C:/Users/Cristina Setimane/Documents/GitHub/rewritejob/backend/app/model.docx')

# Localizar o parágrafo que servirá como modelo
texto_modelo = "Copiar parágrafo"
paragrafo_modelo = localizar_paragrafo_com_texto(doc, texto_modelo)

if paragrafo_modelo is not None:
    # Novos textos a serem adicionados
    novos_textos = [
        "Este é o novo texto 1.",
        "Este é o novo texto 2.",
        "Este é o novo texto 3."
    ]

    for novo_texto in novos_textos:
        # Adicionar um novo parágrafo e copiar estilos do parágrafo modelo
        novo_paragrafo = doc.add_paragraph()
        copiar_estilos(paragrafo_modelo, novo_paragrafo)

        # Alterar o texto do novo parágrafo
        novo_paragrafo.add_run(novo_texto)

    # Remover o parágrafo original
    remover_paragrafo(doc, paragrafo_modelo)

    # Adicionar o sumário automático
    adicionar_sumario(doc)

    # Salvar o documento com as alterações
    doc.save('C:/Users/Cristina Setimane/Documents/GitHub/rewritejob/backend/app/documento_modificado.docx')
else:
    print("Parágrafo com o texto 'Copiar parágrafo' não encontrado.")
