from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# Função para converter <strong> tags para texto em negrito
def convert_html_to_bold(doc, html_text):
    # Separar o texto em partes, identificando onde há <strong> e </strong>
    parts = re.split(r'(<strong>|</strong>)', html_text)
    
    # Criar um novo parágrafo
    paragraph = doc.add_paragraph()
    
    bold = False
    for part in parts:
        if part == '<strong>':
            bold = True  # Ativar negrito
        elif part == '</strong>':
            bold = False  # Desativar negrito
        else:
            run = paragraph.add_run(part)
            run.bold = bold  # Aplicar ou não o negrito
    
    # Justificar o parágrafo
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Criar um novo documento
doc = Document()

# Adicionar o título principal
title = doc.add_heading('Título Principal do Documento', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Exemplo de texto com tags HTML
texto_com_html = """
<strong>Resistores:</strong> São elementos que se opõem à passagem da corrente elétrica, 
convertendo a energia elétrica em calor.
<strong>Capacitores:</strong> Armazenam energia elétrica em um campo elétrico.
<strong>Indutores:</strong> Armazenam energia elétrica em um campo magnético.
"""

# Converter e adicionar o texto formatado
convert_html_to_bold(doc, texto_com_html)

# Adicionar outro parágrafo com um título de nível 2
doc.add_heading('Componentes Passivos', level=2)
paragraph = doc.add_paragraph("""
Os componentes passivos são aqueles que, em condições normais de funcionamento, 
não geram energia elétrica, mas consomem ou armazenam a energia fornecida por outros componentes.
""")
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Aplicar espaçamento entre parágrafos
for paragraph in doc.paragraphs:
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.line_spacing = Pt(12)

# Salvar o documento
doc.save('documento_formatado.docx')
