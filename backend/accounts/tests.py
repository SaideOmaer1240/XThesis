from docx import Document

class DocxEditor:
    def __init__(self, file_path):
        """
        Inicializa o editor de documentos com o caminho do arquivo .docx.

        Args:
            file_path (str): O caminho para o arquivo .docx que será editado.
        """
        self.file_path = file_path
        self.document = Document(file_path)
        self.style_to_apply = None

    def remover_paragrafos_com_palavra(self, palavra_inicial):
        """
        Remove parágrafos que começam com uma palavra específica.

        Args:
            palavra_inicial (str): A palavra que, se encontrada no início de um parágrafo, causará sua remoção.
        """
        novos_paragrafos = []
        i = 0
        while i < len(self.document.paragraphs):
            para = self.document.paragraphs[i]
            
            # Verifica se o parágrafo começa com a palavra específica
            if para.text.startswith(palavra_inicial):
                # Remove o parágrafo que começa com a palavra específica
                pass
            else:
                # Adiciona o parágrafo na lista de novos parágrafos
                novos_paragrafos.append(para.text)
            
            i += 1
        
        # Limpa todos os parágrafos no documento original
        for para in self.document.paragraphs:
            para.clear()
        
        # Recria o documento sem os parágrafos que começam com a palavra específica
        for text in novos_paragrafos:
            self.document.add_paragraph(text)

    def localizar_e_copiar_estilo(self, palavra_localizar):
        """
        Localiza um parágrafo contendo uma palavra específica e copia seu estilo.

        Args:
            palavra_localizar (str): A palavra para localizar o parágrafo a partir do qual o estilo será copiado.
        """
        for para in self.document.paragraphs:
            if palavra_localizar in para.text:
                # Copia o estilo do parágrafo encontrado
                self.style_to_apply = para.style
                break

    def aplicar_estilo(self, palavra_aplicar):
        """
        Localiza um parágrafo contendo uma palavra específica e aplica o estilo copiado anteriormente.

        Args:
            palavra_aplicar (str): A palavra para localizar o parágrafo ao qual o estilo será aplicado.
        """
        if self.style_to_apply is None:
            raise ValueError("Nenhum estilo foi copiado. Execute 'localizar_e_copiar_estilo' primeiro.")

        for para in self.document.paragraphs:
            if palavra_aplicar in para.text:
                # Aplica o estilo copiado ao parágrafo encontrado
                para.style = self.style_to_apply
                break

    def salvar(self):
        """
        Salva o documento editado no mesmo caminho de onde foi carregado.
        """
        self.document.save(self.file_path)

# Exemplo de uso da classe DocxEditor
file_path = 'C:/Users/Cristina Setimane/Documents/GitHub/XThesis/backend/Equação quadrática de segundo grau.docx'
editor = DocxEditor(file_path)

# Remover parágrafos que começam com a palavra 'Aqui está'
editor.remover_paragrafos_com_palavra('Aqui está')

# Localizar e copiar o estilo de um parágrafo contendo a palavra 'EstiloOrigem'
editor.localizar_e_copiar_estilo('EstiloOrigem')

# Aplicar o estilo copiado a um parágrafo contendo a palavra 'EstiloDestino'
editor.aplicar_estilo('EstiloDestino')

# Salvar as alterações no mesmo arquivo
editor.salvar()
