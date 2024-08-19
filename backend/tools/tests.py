from learnfetch import Pesquisador 
from docx import Document
 
researcher = Pesquisador()

# Realizar uma busca
termo_de_busca = input("Topico: ")
resultados = researcher.get_response(termo_de_busca)

    # Função para criar uma ficha de leitura para cada item
def adicionar_ficha_ao_documento(doc, titulo, conteudo):
        doc.add_heading('Ficha de Leitura', level=1)
        doc.add_heading('Título:', level=2)
        doc.add_paragraph(titulo)
        doc.add_heading('Resumo:', level=2)
        doc.add_paragraph(conteudo)
        doc.add_heading('Comentários:', level=2)
        doc.add_paragraph('Adicione aqui suas observações pessoais.')
        doc.add_heading('Questões levantadas:', level=2)
        doc.add_paragraph('Liste aqui as questões ou dúvidas surgidas durante a leitura.')
        doc.add_paragraph("\n" + "="*50 + "\n")

    # Criar um documento Word
documento = Document()

    # Verificar se a chave 'results' está presente no dicionário retornado
if resultados:
        # Iterar sobre cada item na lista de resultados
        for item in resultados:
            titulo = item.get('title', 'Título não encontrado')
            conteudo = item.get('content', 'Conteúdo não encontrado')
            adicionar_ficha_ao_documento(documento, titulo, conteudo)
else:
    print("Nenhum resultado encontrado.")

    # Salvar o documento
nome_arquivo = f'ficha sobre {termo_de_busca}.docx'
documento.save(nome_arquivo)
print(f"As fichas de leitura foram salvas no arquivo {nome_arquivo}.")
