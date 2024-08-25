from docx import Document
import re

class WordFormatter:
    def __init__(self, file_path):
        """
        Inicializa a classe com o caminho do arquivo.
        """
        self.file_path = file_path
        self.doc = Document(file_path)
    
    def replace_strong_with_bold(self):
        """
        Substitui todas as tags <strong> por texto em negrito nos parágrafos do documento.
        """
        for paragraph in self.doc.paragraphs:
            if '<strong>' in paragraph.text:
                self._process_paragraph_bold(paragraph)
    
    def replace_i_with_italic(self):
        """
        Substitui todas as tags <i> por texto em itálico nos parágrafos do documento.
        """
        for paragraph in self.doc.paragraphs:
            if '<i>' in paragraph.text:
                self._process_paragraph_italic(paragraph)
    
    def replace_sub_with_subscript(self):
        """
        Substitui todas as tags <sub> por texto em subscrito nos parágrafos do documento.
        """
        for paragraph in self.doc.paragraphs:
            if '<sub>' in paragraph.text:
                self._process_paragraph_subscript(paragraph)

    def replace_sup_with_superscript(self):
        """
        Substitui todas as tags <sup> por texto em sobrescrito nos parágrafos do documento.
        """
        for paragraph in self.doc.paragraphs:
            if '<sup>' in paragraph.text:
                self._process_paragraph_superscript(paragraph)

    def _process_paragraph_bold(self, paragraph):
        """
        Processa um parágrafo individual para substituir <strong> por negrito.
        """
        # Separar o texto em partes
        parts = re.split(r'(<strong>|</strong>)', paragraph.text)
        
        # Limpar o parágrafo existente
        paragraph.clear()

        bold = False
        for part in parts:
            if part == '<strong>':
                bold = True
            elif part == '</strong>':
                bold = False
            else:
                run = paragraph.add_run(part)
                run.bold = bold
    
    def _process_paragraph_italic(self, paragraph):
        """
        Processa um parágrafo individual para substituir <i> por itálico.
        """
        # Separar o texto em partes
        parts = re.split(r'(<i>|</i>)', paragraph.text)
        
        # Limpar o parágrafo existente
        paragraph.clear()

        italic = False
        for part in parts:
            if part == '<i>':
                italic = True
            elif part == '</i>':
                italic = False
            else:
                run = paragraph.add_run(part)
                run.italic = italic

    def _process_paragraph_subscript(self, paragraph):
        """
        Processa um parágrafo individual para substituir <sub> por subscrito.
        """
        # Separar o texto em partes
        parts = re.split(r'(<sub>|</sub>)', paragraph.text)
        
        # Limpar o parágrafo existente
        paragraph.clear()

        subscript = False
        for part in parts:
            if part == '<sub>':
                subscript = True
            elif part == '</sub>':
                subscript = False
            else:
                run = paragraph.add_run(part)
                run.font.subscript = subscript

    def _process_paragraph_superscript(self, paragraph):
        """
        Processa um parágrafo individual para substituir <sup> por sobrescrito.
        """
        # Separar o texto em partes
        parts = re.split(r'(<sup>|</sup>)', paragraph.text)
        
        # Limpar o parágrafo existente
        paragraph.clear()

        superscript = False
        for part in parts:
            if part == '<sup>':
                superscript = True
            elif part == '</sup>':
                superscript = False
            else:
                run = paragraph.add_run(part)
                run.font.superscript = superscript

    def save(self, output_path=None):
        """
        Salva o documento, sobrescrevendo o original ou criando um novo arquivo.
        """
        if output_path:
            self.doc.save(output_path)
        else:
            self.doc.save(self.file_path)

 

